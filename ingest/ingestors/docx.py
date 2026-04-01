# ingest/ingestors/docx.py — DOCX ingestor.
#
# Text   : Docx2txtLoader — full document text as a single Document
# Tables : python-docx Document.tables — one table chunk per table,
#          formatted as a markdown table
# Images : not extracted (python-docx image extraction requires additional
#          zip parsing; DOCX charts are rarely embedded as raster images)

import os
from typing import List

from docx import Document as DocxDocument
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document

from utils.dataclasses import IngestorResult


def _to_markdown(rows: List[List[str]]) -> str:
    """Convert a list of cell-string rows into a markdown table."""
    if len(rows) < 2:
        return ""
    ncols = max(len(row) for row in rows)

    def pad(row: List[str]) -> List[str]:
        return row + [""] * (ncols - len(row))

    header = pad(rows[0])
    md_rows = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * ncols) + " |",
    ]
    for row in rows[1:]:
        md_rows.append("| " + " | ".join(pad(row)) + " |")

    return "\n".join(md_rows)


class DocxIngestor:
    def ingest(self, path: str) -> IngestorResult:
        source = os.path.basename(path)

        # ── Text ─────────────────────────────────────────────────────────────
        text_docs: List[Document] = Docx2txtLoader(path).load()
        for doc in text_docs:
            doc.metadata["source"] = source

        # ── Tables ───────────────────────────────────────────────────────────
        table_chunks: List[dict] = []
        docx = DocxDocument(path)
        for table in docx.tables:
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            md = _to_markdown(rows)
            if md:
                table_chunks.append({
                    "text": md,
                    "source_file": source,
                    "page_number": 0,   # DOCX has no native page numbers
                })

        print(f"[DocxIngestor] {source}: {len(text_docs)} doc(s), "
              f"{len(table_chunks)} table(s)")

        return IngestorResult(
            text_docs=text_docs,
            table_chunks=table_chunks,
            image_crops=[],
        )
